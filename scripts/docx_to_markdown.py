"""Convert a Word document to Markdown, preserving what actually matters.

Written because `textutil` flattens the policy documents badly: it drops every
table into a run of orphaned lines, loses all bold, and turns numbered lists
into plain paragraphs. The AIP policies carry four tables and 50-70 bold runs
each, so that loss is most of the document's structure.

This walks the document body in order, so tables stay where the author put
them, and reads run-level formatting so emphasis survives.
"""
import re

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(parent):
    """Yield Paragraph and Table objects in document order."""
    from docx.oxml.ns import qn

    for child in parent.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _escape(text):
    """Escape the few Markdown characters that appear in legal prose."""
    return text.replace("|", "\\|")


def _render_run(run):
    """One run, with its emphasis. Whitespace stays outside the markers."""
    text = _escape(run.text)
    if not text:
        return ""
    lead = text[: len(text) - len(text.lstrip())]
    trail = text[len(text.rstrip()) :]
    core = text.strip()
    if core:
        if run.bold:
            core = f"**{core}**"
        if run.italic:
            core = f"_{core}_"
    return f"{lead}{core}{trail}"


def _hyperlink_target(paragraph, element):
    """The URL behind a w:hyperlink, or None for an internal anchor."""
    from docx.oxml.ns import qn

    rel_id = element.get(qn("r:id"))
    if not rel_id:
        return None
    return paragraph.part.rels[rel_id].target_ref


def _runs_to_markdown(paragraph):
    """Render a paragraph, preserving bold, italic and links.

    `paragraph.runs` skips anything inside a w:hyperlink, so reading it alone
    silently drops every link in the document. The policies carry 23 of them
    between them: the contact address in five of the seven, and in the Cookie
    Policy ten links to the privacy notices of the processors it names, which
    is most of the point of a cookie policy. The loss was invisible because
    the link text survived as ordinary prose, except where the whole visible
    text was the link — the Privacy Policy's contact line then ended on a
    dangling em dash with nothing after it.

    So walk the paragraph's children in order instead, and turn each link into
    a Markdown one. Bare mailto: addresses are written as plain text where the
    label already is the address, since a mailto link that reads as its own
    URL gains nothing and the site turns addresses into links anyway.
    """
    from docx.oxml.ns import qn
    from docx.text.run import Run

    out = []
    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:r"):
            out.append(_render_run(Run(child, paragraph)))
        elif child.tag == qn("w:hyperlink"):
            label = "".join(
                _render_run(Run(r, paragraph))
                for r in child.findall(qn("w:r"))
            ).strip()
            if not label:
                continue
            target = _hyperlink_target(paragraph, child)
            if not target or target.startswith("mailto:"):
                out.append(label)
            else:
                out.append(f"[{label}]({target})")
    return "".join(out).strip()


def _is_list(paragraph):
    return paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None


def _callout_to_markdown(table):
    """Render a one-by-one table as a blockquote.

    The policies use a single shaded cell as a callout box — "Our commitment",
    "Secrets rule — absolute", and seven more across the suite. A one-cell
    table is not tabular data and rendering it as a Markdown table produces a
    bordered box with an empty header, which is not what the author drew.

    Worse, `cell.text` flattens the cell's paragraphs into one line and drops
    every run's formatting, so the bold label ran straight into the sentence
    after it: "Our commitment AIP does not sell, rent...". Walking the
    paragraphs keeps the label on its own line and keeps the bold.
    """
    lines = []
    for paragraph in table.rows[0].cells[0].paragraphs:
        text = _runs_to_markdown(paragraph)
        if text:
            lines.append(text)
    if not lines:
        return ""
    return "\n>\n".join(f"> {line}" for line in lines)


def _table_to_markdown(table):
    """Render a Word table as a Markdown table."""
    if len(table.rows) == 1 and len(table.rows[0].cells) == 1:
        return _callout_to_markdown(table)

    rows = []
    for row in table.rows:
        cells = [" ".join(c.text.split()) for c in row.cells]
        # Word merges are expressed as repeated cells; collapse runs of the
        # same value so a merged header does not repeat across the row.
        deduped = [c for i, c in enumerate(cells) if i == 0 or c != cells[i - 1]]
        rows.append([_escape(c) for c in deduped])

    if not rows:
        return ""

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    head, *body = rows
    out = ["| " + " | ".join(head) + " |",
           "| " + " | ".join("---" for _ in head) + " |"]
    out += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(out)


def convert(path, stop_before=None, skip_until=None):
    """Convert `path` to Markdown.

    skip_until — a regex; everything before the first paragraph matching it
                 is dropped (used to cut the header table and approval note).
    stop_before — a regex; conversion stops at the first match.
    """
    doc = Document(path)
    out, started = [], skip_until is None

    for block in _iter_block_items(doc):
        if isinstance(block, Table):
            if started:
                out.append("\n" + _table_to_markdown(block) + "\n")
            continue

        text = _runs_to_markdown(block)
        if not text:
            continue

        plain = re.sub(r"[*_]", "", text)

        if not started:
            if re.match(skip_until, plain):
                started = True
            else:
                continue
        if stop_before and re.match(stop_before, plain):
            break

        # A paragraph can carry no style at all, in which case .style is None.
        style = (getattr(block.style, "name", None) or "").lower()
        if re.match(r"^\d+\.\s+\S", plain) and len(plain) < 90:
            # Section headings are numbered and short; a numbered list item
            # inside a section is not.
            out.append(f"\n## {plain}\n")
        elif style.startswith("heading"):
            level = "##" if style.endswith(("1", "2")) else "###"
            out.append(f"\n{level} {plain}\n")
        elif _is_list(block):
            out.append(f"- {text}")
        else:
            out.append(f"\n{text}\n")

    md = "\n".join(out)
    return re.sub(r"\n{3,}", "\n\n", md).strip() + "\n"


def header_table(path):
    """The key/value table at the top of every AIP policy."""
    doc = Document(path)
    for block in _iter_block_items(doc):
        if isinstance(block, Table):
            pairs = {}
            for row in block.rows:
                cells = [" ".join(c.text.split()) for c in row.cells]
                deduped = [c for i, c in enumerate(cells) if i == 0 or c != cells[i - 1]]
                if len(deduped) >= 2 and deduped[0]:
                    pairs[deduped[0]] = deduped[1]
            if "Policy ID" in pairs:
                return pairs
    return {}
